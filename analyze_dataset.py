import json
import sys
from pathlib import Path

try:
    from docx import Document

    dataset_path = Path(r'c:\Users\anjal\Downloads\[PUB] India_runs_data_and_ai_challenge (1)\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge')

    jd_doc = Document(dataset_path / 'job_description.docx')
    print("=" * 80)
    print("JOB DESCRIPTION")
    print("=" * 80)
    for para in jd_doc.paragraphs:
        print(para.text)
    print()

    signals_doc = Document(dataset_path / 'redrob_signals_doc.docx')
    print("=" * 80)
    print("REDROB SIGNALS DOCUMENT")
    print("=" * 80)
    for para in signals_doc.paragraphs:
        print(para.text)
    print()

    spec_doc = Document(dataset_path / 'submission_spec.docx')
    print("=" * 80)
    print("SUBMISSION SPECIFICATION")
    print("=" * 80)
    for para in spec_doc.paragraphs:
        print(para.text)
    print()

    readme_doc = Document(dataset_path / 'README.docx')
    print("=" * 80)
    print("README")
    print("=" * 80)
    for para in readme_doc.paragraphs:
        print(para.text)
    print()

except ImportError:
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    print("Please run this script again after installation.")
